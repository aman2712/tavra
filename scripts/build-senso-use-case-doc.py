from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "plan" / "tavra-senso-use-case.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
INK = RGBColor(31, 41, 55)
GRAY = RGBColor(91, 101, 116)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
PALE_BLUE = "EEF5FB"
WHITE = RGBColor(255, 255, 255)
TOTAL_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != TOTAL_DXA:
        raise ValueError(f"Table widths must total {TOTAL_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TOTAL_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color="D6DCE5", size="5"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run_element = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run_element.append(text)
    fld.append(run_element)
    paragraph._p.append(fld)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, *, bold_lead=None, italic=False, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, italic=italic)
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.left_indent = Inches(0.5 + (0.25 * level))
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def new_numbering_sequence(doc):
    numbering = doc.part.numbering_part.element
    style_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    source_num = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == style_num_id
    )
    abstract_num_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    next_num_id = max(
        int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))
    ) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return next_num_id


def add_number(doc, text, num_id):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TOTAL_DXA])
    set_table_borders(table, color="BCD4E6", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    title_run = paragraph.add_run(title)
    set_run_font(title_run, color=NAVY, bold=True)
    body_run = paragraph.add_run(f"\n{text}")
    set_run_font(body_run, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TOTAL_DXA])
    set_table_borders(table, color="D4D8DE", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, name="Courier New", size=9.2, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, BLUE_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=9.5, color=NAVY, bold=True)

    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(value)
            set_run_font(run, size=9.3, color=INK)

    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

heading_tokens = {
    1: (16, BLUE, 16, 8),
    2: (13, BLUE, 12, 6),
    3: (12, DARK_BLUE, 8, 4),
}
for level, (size, color, before, after) in heading_tokens.items():
    style = styles[f"Heading {level}"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_style in ("List Bullet", "List Bullet 2", "List Number"):
    style = styles[list_style]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.167

header = section.header.paragraphs[0]
header.paragraph_format.space_after = Pt(0)
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_run = header.add_run("TAVRA  |  TEAM RECOVERY CONTEXT DESIGN")
set_run_font(header_run, size=8.5, color=GRAY, bold=True)

footer = section.footer.paragraphs[0]
footer.paragraph_format.space_before = Pt(0)
footer.paragraph_format.right_indent = Inches(0.12)
add_page_number(footer)

# Memo masthead
kicker = doc.add_paragraph()
kicker.paragraph_format.space_before = Pt(8)
kicker.paragraph_format.space_after = Pt(3)
kicker_run = kicker.add_run("PRODUCT & KNOWLEDGE DESIGN")
set_run_font(kicker_run, size=10, color=BLUE, bold=True)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(5)
title.paragraph_format.keep_with_next = True
title_run = title.add_run("Tavra × Senso")
set_run_font(title_run, size=26, color=NAVY, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
subtitle.paragraph_format.keep_with_next = True
subtitle_run = subtitle.add_run("Identity-aware Team Recovery with verified context and delta-only questions")
set_run_font(subtitle_run, size=13.5, color=GRAY)

metadata = [
    ("Status", "Locked use case for the hackathon vertical slice"),
    ("Scope", "Team Recovery only; no Personal Recovery fallback"),
    ("Primary scenario", "Delayed baggage before an 08:00 client meeting in Boston"),
    ("Prepared", "1 August 2026"),
]
for label, value in metadata:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=NAVY)
    value_run = paragraph.add_run(value)
    set_run_font(value_run)

doc.add_paragraph().paragraph_format.space_after = Pt(5)

add_callout(
    doc,
    "Decision",
    "Tavra resolves the Linq sender number to an exact company employee, retrieves that employee's profile and permitted policy context from Senso, and asks only for missing, stale, conflicting, or checkout-critical confirmations. Senso supplies evidence and outcome memory; deterministic Tavra code controls identity, policy calculations, authorization, and money.",
)

add_heading(doc, "1. Product scope", 1)
add_body(
    doc,
    "The hackathon product is an organization-backed Team Recovery workflow. The employee states the disruption once. Tavra reconstructs what the company already knows—identity, sizes, employee category, spending authority, evidence rules and prior outcomes—then requests only the remaining information needed to recover the trip safely.",
)
add_bullet(doc, "In scope: employee identification, stored sizes and preferences, company policy, airline/airport evidence collection, merchant selection, exact approval, claim-ready evidence and verified outcome memory.")
add_bullet(doc, "Out of scope: Personal Recovery, semantic identity guessing, unrestricted internet research, autonomous policy exceptions, card data in model context, and submitting reimbursement claims.")
add_bullet(doc, "Unknown company phone numbers receive no employee information and no personal-mode fallback; they are routed to an organization-verification path.")

add_heading(doc, "2. Responsibility model", 1)
add_table(
    doc,
    ["Layer", "Owns", "Must not own"],
    [
        ("Linq", "Inbound sender identity and iMessage delivery", "Employee policy or payment authority"),
        ("Tavra backend", "Exact identity mapping, state, policy calculation, authorization and audit", "Semantic answers or card credentials"),
        ("OpenAI", "Incident extraction, missing-field reasoning and concise explanations", "Identity lookup, final spend decision or invented facts"),
        ("Senso", "Scoped employee/company context, cited evidence and verified outcome memory", "Identity authentication, live truth or final payment decision"),
        ("Prava", "Passkey approval, scoped credentials and payment status", "Product eligibility or delivery judgement"),
    ],
    [1500, 3990, 3870],
)

add_heading(doc, "3. Identity-aware request sequence", 1)
add_code_block(
    doc,
    "iMessage → Linq sender handle → exact Tavra employee lookup\n"
    "         → employee/profile/policy content IDs\n"
    "         → OpenAI incident extraction\n"
    "         → scoped Senso context retrieval\n"
    "         → deterministic missing-field + policy evaluation\n"
    "         → one concise confirmation message in iMessage",
)
identity_sequence = new_numbering_sequence(doc)
add_number(doc, "Normalize the Linq sender number to E.164 and resolve it through a private exact mapping to company ID and employee ID.", identity_sequence)
add_number(doc, "Fetch the exact employee profile document and query only the permitted company policy content IDs. Never search all employees semantically.", identity_sequence)
add_number(doc, "Extract incident, business objective, deadline and destination from the message; treat omitted airline, airport and baggage reference as unknown.", identity_sequence)
add_number(doc, "Build a field matrix of known, missing, stale, conflicting and checkout-critical values.", identity_sequence)
add_number(doc, "Return known values, explain the applicable budget, and ask one bundled delta-only question for the remaining material fields.", identity_sequence)

add_heading(doc, "4. Employee and policy knowledge model", 1)
add_table(
    doc,
    ["Domain", "Example facts", "Source", "Decision use"],
    [
        ("Employee profile", "T-shirt M; waist 32; inseam missing; preferred fit", "HR sync or employee-confirmed profile", "Pre-fill variants and ask only for missing confirmation"),
        ("Employee category", "Client-facing traveller", "Company directory", "Select the applicable policy and approval route"),
        ("Company policy", "$175 incident ceiling; eligible categories", "Company-owned policy", "Deterministic ALLOW / REAUTHORIZE / DENY"),
        ("External policy", "Airline evidence and timing requirements", "Official runtime web snapshot", "Prepare claim-ready evidence and expose gaps"),
        ("Merchant/product", "Price, variant, condition, delivery and returns", "Official live evidence or synthetic sandbox source", "Rank, reject and reconcile candidates"),
        ("Outcomes", "Quote drift, item match and checkout status", "Observed Tavra transaction", "Improve later merchant decisions without overstating delivery"),
    ],
    [1500, 2490, 2280, 3090],
)

add_heading(doc, "Profile field states", 2)
add_bullet(doc, "Known and current: report the value; confirm again only when it becomes an approved item variant.")
add_bullet(doc, "Known but stale: report the stored value and request confirmation.")
add_bullet(doc, "Missing: explicitly say company data does not include it and ask for the value.")
add_bullet(doc, "Conflicting: show the conflict without selecting a value.")
add_bullet(doc, "Updated: write back only after explicit employee confirmation, with source and timestamp. Never infer a profile update from an order.")

add_heading(doc, "5. Delta-only conversation behavior", 1)
add_body(doc, "Employee message:", bold_lead="Employee message:")
add_callout(doc, "Inbound", "My bag is delayed. I have a client meeting at 8am in Boston.", fill="F7F8FA")
add_body(doc, "When an inseam is missing:", bold_lead="When an inseam is missing:")
add_callout(
    doc,
    "Tavra",
    "Sorry about the delay. I have a medium T-shirt and 32-inch trouser waist on file, but no inseam. Your policy allows up to $175 for essential clothing and toiletries. What inseam should I use, and which airline and arrival airport were involved?",
)
add_body(doc, "When both sizes exist:", bold_lead="When both sizes exist:")
add_callout(
    doc,
    "Tavra",
    "Sorry about the delay. I have medium T-shirts and 32×30 trousers on file, and your policy allows up to $175 for essential clothing and toiletries. Are those sizes still correct, and which airline and arrival airport were involved?",
)
add_body(
    doc,
    "Boston is a destination, not proof of the arrival airport. Tavra must not infer BOS, an airline, or reimbursement rules. If itinerary data is not connected, it asks for the missing facts.",
)

add_heading(doc, "6. Policy and budget enforcement", 1)
add_body(
    doc,
    "Senso supplies the employee category and policy facts. Tavra code performs arithmetic and policy enforcement. The model may explain the result but cannot alter the ceiling or approval route.",
)
add_code_block(
    doc,
    "employee_category: client_facing_traveller\n"
    "incident_allowance: USD 175 inclusive\n"
    "self_approval: total <= 175 and all rules pass\n"
    "manager_reauthorization: 175 < total <= 300\n"
    "deny: total > 300 or prohibited term change",
)
add_bullet(doc, "All taxes, delivery charges and fees count toward the ceiling.")
add_bullet(doc, "Exact size, condition, quantity, merchant, delivery deadline and maximum total become part of the Recovery Authorization.")
add_bullet(doc, "Any material checkout change invalidates the old authorization and produces REAUTHORIZE or DENY.")

add_heading(doc, "7. Static demo corpus versus runtime evidence", 1)
add_table(
    doc,
    ["Generated before demo", "Retrieved during workflow", "Boundary"],
    [
        ("Synthetic employee profiles and identity linkage", "Optional authoritative HR/profile refresh", "Identity resolution remains exact and private"),
        ("Synthetic company recovery and evidence policies", "Real company policy only when supplied by the organization", "Never present generated policy as a real employer rule"),
        ("Fictional merchant A/B and product catalog", "Real merchant price, inventory, delivery, returns and seller", "Synthetic evidence is sandbox-only"),
        ("Prior synthetic merchant outcome", "Verified outcome from the completed checkout", "Checkout does not prove delivery"),
        ("Airline and merchant snapshot templates", "Official airline, regulator, airport and merchant pages", "Allowlist, timestamp and hash every external source"),
    ],
    [2700, 3390, 3270],
)
add_body(
    doc,
    "For live evidence, Tavra fetches an allowlisted official page, records provenance and a content hash, fills a structured snapshot, ingests it into Senso, waits for compilation, and queries only that snapshot. If freshness or provenance fails, the fact remains unknown. The MVP prepares a claim-ready Recovery Receipt; it does not claim to submit reimbursement.",
)

add_heading(doc, "8. Senso outcome memory", 1)
add_code_block(
    doc,
    "Verified checkout → Merchant Outcome Record → Senso Outcomes folder\n"
    "                  → compilation complete → next scoped query\n"
    "                  → prior evidence influences the next merchant decision",
)
add_bullet(doc, "Store quoted and final totals, amount drift, approved-item match, add-ons, checkout state, order-ID presence and verification time.")
add_bullet(doc, "Keep delivery status `not_verified` until an actual delivery event is observed.")
add_bullet(doc, "Carry a single correlation ID through Linq events, Senso queries, OpenAI traces, Prava sessions, checkout records and the Recovery Receipt.")

add_heading(doc, "9. Hackathon corpus delivered with this brief", 1)
add_table(
    doc,
    ["Package", "Purpose", "Trust"],
    [
        ("employees/emp_demo_001.md", "Profile with known T-shirt/waist and missing inseam", "Synthetic company authority"),
        ("policies/*.md", "Budget, categories, approval and evidence rules", "Synthetic company authority"),
        ("merchants/*.md + products/*.json", "Cheaper rejected option and trusted selected option", "Synthetic sandbox evidence"),
        ("outcomes/*.json", "Prior verified checkout-shaped memory", "Synthetic outcome for retrieval demo"),
        ("templates/*", "Runtime airline, merchant and outcome records", "Not evidence until populated"),
        ("queries/retrieval-playbook.md", "Scoped prompts and delta-only behavior", "Application-owned instructions"),
        ("demo-config/identity-map.example.json", "Exact phone-to-employee mapping placeholder", "Private backend configuration"),
    ],
    [2700, 4010, 2650],
)
add_body(doc, "Repository location: `senso/demo-corpus/` with private demo identity configuration under `senso/demo-config/`.")

add_heading(doc, "10. Security and privacy controls", 1)
add_bullet(doc, "Do not place raw phone numbers in Senso employee documents. Map phone to employee/profile content ID in Tavra.")
add_bullet(doc, "Restrict retrieval using the exact employee profile content ID and permitted policy IDs; never issue an organization-wide employee query.")
add_bullet(doc, "Use separate Senso credentials: viewer access for retrieval and editor access only for controlled outcome/profile writes.")
add_bullet(doc, "Do not expose one employee's profile, policy attachments or outcome history to another employee.")
add_bullet(doc, "Do not store card credentials, one-time payment credentials, passwords, medical details or unnecessary personal data in Senso or OpenAI context.")
add_bullet(doc, "Log source IDs and decisions, not message bodies or personal profile values.")

add_heading(doc, "11. Demo acceptance criteria", 1)
add_table(
    doc,
    ["Proof", "Required visible behavior"],
    [
        ("Exact identity", "Known Linq sender resolves to one employee; unknown number reveals no profile"),
        ("Profile reuse", "Tavra reports stored size and asks only for the missing inseam"),
        ("Policy grounding", "The $175 limit and eligible categories come from the correct company policy"),
        ("Context changes action", "Cheaper Merchant A is rejected because evidence is insufficient"),
        ("Safe approval", "A changed amount or term triggers reauthorization or denial"),
        ("Outcome memory", "The Merchant Outcome Record is written, compiled and returned by a second query"),
        ("Truthful reimbursement", "Receipt is claim-ready; no claim-submission or delivery claim is fabricated"),
    ],
    [2500, 6860],
)

add_heading(doc, "12. Recommended implementation order", 1)
implementation_sequence = new_numbering_sequence(doc)
add_number(doc, "Create the Tavra identity map and exact employee resolver using the Linq sender handle.", implementation_sequence)
add_number(doc, "Create Senso folders, ingest the generated seed corpus and capture each content ID.", implementation_sequence)
add_number(doc, "Implement exact profile retrieval and policy queries restricted to known content IDs.", implementation_sequence)
add_number(doc, "Add the deterministic field-state and delta-only clarification builder.", implementation_sequence)
add_number(doc, "Add allowlisted official-source retrieval and runtime Senso snapshot ingestion.", implementation_sequence)
add_number(doc, "Add merchant candidate evaluation, Recovery Authorization and Prava checkout guardrails.", implementation_sequence)
add_number(doc, "Write the verified Merchant Outcome Record, wait for compilation and demonstrate the second query.", implementation_sequence)

add_heading(doc, "Official platform references", 1)
references = [
    ("Senso introduction", "https://docs.senso.ai/docs/introduction"),
    ("Senso core concepts: ingest, compile and query", "https://docs.senso.ai/docs/concepts"),
    ("Senso knowledge base", "https://docs.senso.ai/docs/knowledge-base"),
    ("Senso permissions and API-key scoping", "https://docs.senso.ai/docs/permissions"),
    ("Tavra project plan", "plan/plan.docx"),
]
for label, url in references:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(6)
    if url.startswith("http"):
        add_hyperlink(paragraph, label, url)
    else:
        run = paragraph.add_run(f"{label}: {url}")
        set_run_font(run)

properties = doc.core_properties
properties.title = "Tavra × Senso — Team Recovery Context and Knowledge Design"
properties.subject = "Identity-aware Team Recovery use case and Senso demo corpus"
properties.author = "Tavra"
properties.keywords = "Tavra, Senso, Team Recovery, Linq, OpenAI, Prava"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
